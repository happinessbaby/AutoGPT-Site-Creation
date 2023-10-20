import os
from langchain.chat_models import ChatOpenAI
from langchain.llms import OpenAI
from langchain.embeddings import OpenAIEmbeddings
from langchain.prompts import ChatPromptTemplate
from langchain.output_parsers import CommaSeparatedListOutputParser
from langchain.prompts import PromptTemplate
from langchain_experimental.smart_llm import SmartLLMChain
from langchain.agents import AgentType, Tool, initialize_agent, create_json_agent
from basic_utils import read_txt
from common_utils import (get_web_resources, retrieve_from_db,  get_generated_responses,calculate_graduation_years, find_resume_content, extract_posting_keywords,
                           extract_resume_fields, search_related_samples, create_sample_tools, extract_personal_information)
from langchain_utils import create_search_tools, create_mapreduce_chain, create_summary_chain, generate_multifunction_response, create_refine_chain, handle_tool_error
from pathlib import Path
import json
from json import JSONDecodeError
from multiprocessing import Process, Queue, Value
from langchain.tools.json.tool import JsonSpec
from langchain.agents.agent_toolkits import JsonToolkit
from typing import Dict, List, Optional
from langchain.document_loaders import BSHTMLLoader
from langchain.tools import tool
from langchain.agents.agent_toolkits import FileManagementToolkit
import docx
import uuid
from docxtpl import DocxTemplate	
from docx import Document
from docx.shared import Inches


from dotenv import load_dotenv, find_dotenv
_ = load_dotenv(find_dotenv()) # read local .env file
# resume_evaluation_path = os.environ["RESUME_EVALUATION_PATH"]
resume_samples_path = os.environ["RESUME_SAMPLES_PATH"]
# TODO: caching and serialization of llm
llm = ChatOpenAI(temperature=0.0)
# llm = OpenAI(temperature=0, top_p=0.2, presence_penalty=0.4, frequency_penalty=0.2)
embeddings = OpenAIEmbeddings()
# TODO: save these delimiters in json file to be loaded from .env
delimiter = "####"
delimiter1 = "````"
delimiter2 = "////"
delimiter3 = "<<<<"
delimiter4 = "****"

# personal_info = ["Name", "Phone", "Email", "LinkedIn", "Website", "JobTitle"]
document = Document()
document.add_heading('Resume Evaluation', 0)


def evaluate_resume(about_me="", resume_file = "", posting_path="") -> str:

    dirname, fname = os.path.split(resume_file)
    filename = Path(fname).stem 
    docx_filename = filename + "_evaluation"+".docx"
    # get resume info
    resume_content = read_txt(resume_file)
    info_dict=get_generated_responses(resume_content=resume_content, posting_path=posting_path, about_me=about_me)
    work_experience_level = info_dict.get("work experience level", "")
    graduation_year = info_dict.get("graduation year", -1)
    years_since_graduation = calculate_graduation_years(graduation_year)
    degree = info_dict.get("degree", -1)
    study = info_dict.get("study", -1)
    job = info_dict.get("job", -1)
    company = info_dict.get("company", -1)
    field_names = info_dict.get("field names", "")

    if work_experience_level=="no experience" or work_experience_level=="entry level" and years_since_graduation<2:
        resume_type = "student"
    elif work_experience_level=="no experience" or work_experience_level=="entry level" and years_since_graduation>=2:
        resume_type =  "functional"
    else:
        resume_type = "chronological"

  #TODO: This query should make suggestions on what type of resume they should write and how to improve the overall impression
    query_overall = f"""Your task is to provide an assessment of a resume delimited by {delimiter} characters.

    resume: {delimiter}{resume_content}{delimiter} \n

    The applicant's work experience level as a {job} is {work_experience_level}.

    Furthermore, it has been {years_since_graduation} years since the applicant graduated with a highest level of education {degree} in {study}. 

    Look at the work experience and skills sections of the resume, if available, and assess if it is written as a  {resume_type} resume. 

    For a student resume type, the focus should be on education, voluntary/work experience, and any skills.

    For a functional resume type, the focus should be on skills, trainings, projects and accomplishments. 

    For a chronological resume type, the focus should be on work experience and any reward, honors, and achievements received on the job.

    """
    tools = create_search_tools("google", 3)
    response = generate_multifunction_response(query_overall, tools)
    # prompt = PromptTemplate.from_template(query_overall)
    # chain = SmartLLMChain(llm=llm, prompt=prompt, n_ideas=3, verbose=True)
    # response = chain.run({})

    # document.add_heading(f"Overall Asessment", level=1)
    # document.add_paragraph(response)
    # document.add_page_break()
    # document.save(docx_filename)

       

    # write_to_docx_template(doc, personal_info, personal_info_dict, docx_filename)

    # Note: document comparison benefits from a clear and simple prompt
    # related_samples = search_related_samples(job, resume_samples_path)
    # sample_tools, tool_names = create_sample_tools(related_samples, "resume")

    # # # process all fields in parallel
    # processes = [Process(target = evaluate_resume_fields, args = (info_dict, field, info_dict.get(field, ""),  sample_tools)) for field in field_names]
    # for p in processes:
    #    p.start()
    # for p in processes:
    #    p.join()


    # return f""" file_path: {docx_filename} """



def evaluate_resume_fields(generated_response: Dict[str, str], field: str, field_content: str, tools: List[Tool]) -> None:

    print(f"CURRENT FIELD IS: {field}")
    if field_content!="":
        job = generated_response.get("job", "")
        company_description = generated_response.get("company description", "")
        job_specification = generated_response.get("job specification", "")
        job_description = generated_response.get("job description", "")
        highest_education_level = generated_response.get("highest education level", "")
        work_experience_level = generated_response.get("work experience level", "")
        # education_level = generated_response.get("education", "")

        advice_query =  f"""how to make resume field {field} ATS-friendly? No formatting advices."""
        advice1 = retrieve_from_db(advice_query)
        advice_query = f"""what to include in {field} of resume for {highest_education_level} and {work_experience_level} as a {job}"""
        advice2 = retrieve_from_db(advice_query)

        query_evaluation = f"""  You are an expert resume field advisor. 

        Generate a list of missing, irrelevant, and not ATS-friendly information in the resume field content. 
        
        Remember to use either job specification or general job description as your guideline along with the expert advice.

        field name: {field}

        field content: {field_content}\n

        job specification: {job_specification}\n

        general job description: {job_description} \n

        expert advice: {advice2} + "\n" + {advice1}

        Your answer should be detailed and only from the field content. Please also provide your reasoning too as in the following examples:

                Missing or Irrelevant Field Content for Work Experience:

                1. Quantative achievement is missing: no measurable metrics or KPIs to highlight any past achievements. 

                2. Front desk receptionist is irrelevant: Experience as a front desk receptionist is not directly related to the role of a data analyst

                3. Date formatting is not ATS-friendly: an ATS-friendly way to write dates is for example, 01/2001 or January 2001

        The above is just an example for your reference. Do not let it be your answer. 
        
        Please ignore all formatting advices as formatting should not be part of the assessment.

        Use your tools if you need to reference other resume.

        """

        evaluation = generate_multifunction_response(query_evaluation, tools)

        with open(f"{field}_evaluation.txt", "x") as f:
           f.write(evaluation)



def reformat_functional_resume(resume_file:str, posting_path:str) -> str:

    resume_content = read_txt(resume_file)
    functional_doc_template = DocxTemplate("./resume_templates/functional.docx")
    info_dict = get_generated_responses(resume_content=resume_content, posting_path=posting_path)
    func = lambda key, default: default if info_dict[key]==-1 else info_dict[key]
    personal_context = {
        "NAME": func("name", "YOUR NAME"),
        "ADDRESS": func("address", "YOUR ADDRESS"),
        "PHONE": func("phone", "YOUR PHONE"),
        "EMAIL": func("email", "YOUR EMAIL"),
        "LINKEDIN": func("linkedin", "YOUR LINKEDIN URL"),
        "WEBSITE": func("website", "WEBSITE"),
    }
    #TODO: save the context dictionary somewhere
    context_keys = ["SKILLS_SUMMARY", "WORK_HISTORY", "PROFESSIONAL_ACCOMPLISHMENTS", "EDUCATION"]
    info_dict_keys = ["summary or objective", "work experience", "professional accomplishment", "education"]
    context_dict = dict(zip(context_keys, info_dict_keys))
    context = {key: None for key in context_keys}
    for key, value in context_dict.items():
        # content = find_resume_content(key, resume_content)
        content = info_dict.get(value, "")
        if key == "SKILLS_SUMMARY":
            job_description = info_dict.get("job description", "")
            job_specification = info_dict.get("job specification", "")
            query = f""" Your task is to improve or write the objective section of the functional resume.

            If you are provided with an existing objective section, use it as your context and build on top of it.
              
            Otherwise, refer to the job specification or job description below, whichever is available.

            objective section: {content}

            job description: {job_description}

            job specification: {job_specification}

            Please write in fewer than five sentences the objective section of the functiional resume using the information above. 

            Incorporate relevant soft skills and hard skills into the objective section. 
            
            """
            tools = create_search_tools("google", 1)
            content = generate_multifunction_response(query, tools)
        elif key=="PROFESSIONAL_ACCOMPLISHMENTS":
            if (Path(posting_path).is_file()):
                keywords = extract_posting_keywords(read_txt(posting_path))
                query = f"""Your task is to catgeorize the professional accomplishments, skills, trainings, projects in the resume delimited with {delimiter} characters under certain skills. 

                Please in total pick at least 3 skill from the following available skillset. 

                skillset: {keywords}.

                Categorize the content into the different skills. For example, your output should look like something like this:

                Skill1:

                    - Examples of projects or situations that utilized this skill
                    - Measurable results and accomplishments

                content: {delimiter}{content}{delimiter}

                """
                tools = create_search_tools("google", 1)
                content = generate_multifunction_response(query, tools)
        context[key] = content
    context.update(personal_context)
    # print(context)
    functional_doc_template.render(context)
    functional_doc_template.save("./test_functional_reformatter.docx")   



def reformat_chronological_resume(resume_file:str, posting_path:str,) -> str:

    resume_content = read_txt(resume_file)
    chronological_resume_template = DocxTemplate("./resume_templates/chronological.docx")
    info_dict = get_generated_responses(resume_content=resume_content, posting_path=posting_path)
    func = lambda key, default: default if info_dict[key]==-1 else info_dict[key]
    personal_context = {
        "NAME": func("name", "YOUR NAME"),
        "ADDRESS": func("address", "YOUR ADDRESS"),
        "PHONE": func("phone", "YOUR PHONE"),
        "EMAIL": func("email", "YOUR EMAIL"),
        "LINKEDIN": func("linkedin", "YOUR LINKEDIN URL"),
        "WEBSITE": func("website", "WEBSITE"),
    }
    # TODO: add awards and honors or professional accomplishments
    context_keys = ["SUMMARY", "PROFESSIONAL_EXPERIENCE", "RELEVANT_SKILLS", "EDUCATION"]
    info_dict_keys = ["summary or objective", "work experience", "skills", "education"]
    context_dict = dict(zip(context_keys, info_dict_keys))
    context = {key: None for key in context_keys}
    for key, value in context_dict.items():
        # content = find_resume_content(key, resume_content)
        content = info_dict.get(value, "")
        if key == "SUMMARY":
            work_experience = info_dict.get("work experience", "")
            query = f""" Your task is to improve or rewrite the summary section of a chronological resume.

            If you are provided with an existing summary section, use it as your context and build on top of it, if needed.
              
            Otherwise, refer to the work experience, if available. 

            summary section: {content}

            work experience: {work_experience}

            Please write in fewer than five sentences the summary section of the chronological resume with the information above.

            If the summary already is already filled with relevant work experience, you can output the original summary section. 
            
            Otherwise, incorporate relevant work experience into the summary section. 

            """
            tools = create_search_tools("google", 1)
            content = generate_multifunction_response(query, tools)
        elif key=="RELEVANT_SKILLS":
            if (Path(posting_path).is_file()):
                keywords = extract_posting_keywords(read_txt(posting_path))
            job_description = info_dict.get("job description", "")
            job_specification = info_dict.get("job specification", "") 
            skills = info_dict.get("skills", "")
            query = f""" 

            Your tasks is to improve the Skills section of the resume. You are provided with a job description or job specificaiton, whichever is available.

            If you are provided with an existing Skills section, use it as your context and build on top of it, if needed.

            You are also provided with a list of important keywords that are in the job posting. Some of them should be included also. 

            skills section: {skills}

            job description: {job_description}
            
            job specification: {job_specification}

            important keywords: {keywords}

            If the skills section exist, add to it relevant skills and remove from it irrelevant skills.

            Otherwise, if the skills section is already well-written, output the original skills section. 

            """
        context[key] = content
    context.update(personal_context)
    # print(context)
    chronological_resume_template.render(context)
    chronological_resume_template.save("./test_chronological_reformatter.docx")   


def reformat_student_resume(resume_file:str, posting_path:str) -> str:

    resume_content = read_txt(resume_file)
    chronological_resume_template = DocxTemplate("./resume_templates/student.docx")
    info_dict = get_generated_responses(resume_content=resume_content, posting_path=posting_path)
    func = lambda key, default: default if info_dict[key]==-1 else info_dict[key]
    personal_context = {
        "NAME": func("name", "YOUR NAME"),
        "ADDRESS": func("address", "YOUR ADDRESS"),
        "PHONE": func("phone", "YOUR PHONE"),
        "EMAIL": func("email", "YOUR EMAIL"),
        "LINKEDIN": func("linkedin", "YOUR LINKEDIN URL"),
        "WEBSITE": func("website", "WEBSITE"),
    }
    #TODO: add volunteer experience
    context_keys = ["OBJECTIVE", "EDUCATION", "AWARDS_HONORS", "SKILLS", "WORK_EXPERIENCE"]
    info_dict_keys = ["summary or objective", "education", "awards and honors", "skills", "work experience"]
    context_dict = dict(zip(context_keys, info_dict_keys))
    context = {key: None for key in context_keys}
    for key, value in context_dict.items():
        # content = find_resume_content(key, resume_content)
        content = info_dict.get(value, "")
    context.update(personal_context)
    # print(context)
    chronological_resume_template.render(context)
    chronological_resume_template.save("./test_student_reformatter.docx")   


# @tool("resume evaluator")
# def resume_evaluator_tool(resume_file: str, job: Optional[str]="", company: Optional[str]="", job_post_link: Optional[str]="") -> str:

#    """Evaluate a resume when provided with a resume file, job, company, and/or job post link.
#         Note only the resume file is necessary. The rest are optional.
#         Use this tool more than any other tool when user asks to evaluate, review, help with a resume. """

#    return evaluate_resume(my_job_title=job, company=company, resume_file=resume_file, posting_path=job_post_link)
      


@tool(return_direct=True)
def resume_evaluator(json_request: str)-> str:

    """Helps to evaluate a resume. Use this tool more than any other tool when user asks to evaluate or review a resume. 

    Input should be  a single string strictly in the following JSON format:  '{{"about_me":"<about_me>", "resume_file":"<resume_file>", "job_post_file":"<job_post_file>"}}' \n

    Leave value blank if there's no information provided. DO NOT MAKE STUFF UP. 

     (remember to respond with a markdown code snippet of a JSON blob with a single action, and NOTHING else) \n

     Output should be using the "get download link" tool in the following single string JSON format: '{{"file_path":"<file_path>"}}'
   
    """

    try:
      json_request = json_request.strip("'<>() ").replace('\'', '\"')
      args = json.loads(json_request)
    except JSONDecodeError as e:
      print(f"JSON DECODE ERROR: {e}")
      return "Format in a single string JSON and try again."

    # if resume doesn't exist, ask for resume
    if ("resume_file" not in args or args["resume_file"]=="" or args["resume_file"]=="<resume_file>"):
      return "Can you provide your resume so I can further assist you? "
    else:
      # may need to clean up the path first
        resume_file = args["resume_file"]
    if ("about_me" not in args or args["about_me"] == "" or args["about_me"]=="<about_me>"):
        about_me = ""
    else:
        about_me = args["about_me"]
    if ("job_post_file" not in args or args["job_post_file"]=="" or args["job_post_file"]=="<job_post_file>"):
        posting_path = ""
    else:
        posting_path = args["job_post_file"]

    return evaluate_resume(about_me=about_me, resume_file=resume_file, posting_path=posting_path)



def processing_resume(json_request: str) -> str:

    """ Input parser: input is LLM's action_input in JSON format. This function then processes the JSON data and feeds them to the resume evaluator. """

    try:
      json_request = json_request.strip("'<>() ").replace('\'', '\"')
      args = json.loads(json_request)
    except JSONDecodeError as e:
      print(f"JSON DECODER ERROR: {e}")
      return "Format in JSON and try again."

    # if resume doesn't exist, ask for resume
    if ("resume_file" not in args or args["resume_file"]=="" or args["resume_file"]=="<resume_file>"):
      return "Stop using the resume evaluator tool. Ask user for their resume."
    else:
      # may need to clean up the path first
        resume_file = args["resume_file"]
    if ("about_me" not in args or args["about_me"] == "" or args["about_me"]=="<about_me>"):
        about_me = ""
    else:
        about_me = args["about_me"]
    if ("job_post_file" not in args or args["job_post_file"]=="" or args["job_post_file"]=="<job_post_file>"):
        posting_path = ""
    else:
        posting_path = args["job_post_file"]
        
    return evaluate_resume(about_me=about_me, resume_file=resume_file, posting_path=posting_path)



def processing_resume2(json_request: str) -> str:

    """ Input parser: input is LLM's action_input in JSON format. This function then processes the JSON data and feeds them to the resume reformatter. """

    try:
      json_request = json_request.strip("'<>() ").replace('\'', '\"')
      args = json.loads(json_request)
    except JSONDecodeError as e:
      print(f"JSON DECODER ERROR: {e}")
      return "Format in JSON and try again."

    # if resume doesn't exist, ask for resume
    if ("resume_file" not in args or args["resume_file"]=="" or args["resume_file"]=="<resume_file>"):
      return "Stop using the resume evaluator tool. Ask user for their resume."
    else:
      # may need to clean up the path first
        resume_file = args["resume_file"]
    if ("resume_type" not in args or args["resume_type"]=="" or args["resume_type"]=="<resume_type>"):
      return "Stop using the resume reformatter tool. Ask user to specify a template or type"
    else:
      # may need to clean up the path first
        resume_type = args["resume_type"]
    if ("job_post_file" not in args or args["job_post_file"]=="" or args["job_post_file"]=="<job_post_file>"):
        posting_path = ""
    else:
        posting_path = args["job_post_file"]
    
    if resume_type == "functional":
        return reformat_functional_resume(resume_file=resume_file, posting_path=posting_path)



def create_resume_evaluator_tool() -> List[Tool]:

    """ Input parser: input is user's input as a string of text. This function takes in text and parses it into JSON format. 
    
    Then it calls the processing_resume function to process the JSON data. """

    name = "resume_evaluator"
    parameters = '{{"about_me":"<about_me>", "resume_file":"<resume_file>", "job_post_file":"<job_post_file>"}}' 
    output = '{{"file_path":"<file_path>"}}'
    description = f"""Evaluate a resume. Use this tool more than any other tool when user asks to evaluate or improves a resume. 
    Do not use this tool is asked to customize or tailr the resume. You should use the "resume_customize_writer" instead.
    Input should be a single string strictly in the following JSON format: {parameters} \n
     Leave value blank if there's no information provided. DO NOT MAKE STUFF UP. 
     (remember to respond with a markdown code snippet of a JSON blob with a single action, and NOTHING else) \n
     Output should be using the "get download link" tool in the following single string JSON format: {output}
    """
    tools = [
        Tool(
        name = name,
        func = processing_resume,
        description = description,
        verbose = False,
        handle_tool_error=True,

        )
    ]
    print("Succesfully created resume evaluator tool.")
    return tools

def create_resume_reformatting_tool() -> List[Tool]:

    name = "resume_formatter"
    parameters = '{{"resume_file":"<resume_file>", "resume_type":"<resume_type>", "job_post_file":"<job_post_file>"}}'
    output = '{{"file_path":"<file_path>"}}'
    description = f""" Reformats a resume. Use this tool more than any other tool when user asks to reformat their resume according to a particular type or template.
    Do not use this tool to evaluate or customize and tailor resume content. This tool should only be used for formatting resume to a particular style.
    Input should be a single string strictly in the followiwng JSON format: {parameters} \n
    Leave value blank if there's no information provided. DO NOT MAKE STUFF UP. 
     (remember to respond with a markdown code snippet of a JSON blob with a single action, and NOTHING else) \n
     Output should be using the "get download link" tool in the following single string JSON format: {output}
    """
    tools = [
        Tool(
        name = name,
        func = processing_resume2,
        description = description,
        verbose = False,
        handle_tool_error=True,

        )
    ]
    print("Succesfully created resume evaluator tool.")
    return tools





if __name__ == '__main__':
    # my_job_title = 'Data Analyst'
    # my_resume_file = './resume_samples/resume2023v3.txt'
    job_posting = "./uploads/file/data_analyst_SC.txt"
    # company = "Southern Company"
    # evaluate_resume(my_job_title =my_job_title, company = company, resume_file=my_resume_file, posting_path = job_posting)
    my_resume_file = "./resume_samples/resume2023v3.txt"
    # reformat_functional_resume(resume_file=my_resume_file, posting_path=job_posting)
    reformat_chronological_resume(resume_file="./resume_samples/sample1.txt", posting_path="./uploads/link/accountant.txt")
    # evaluate_resume(resume_file=my_resume_file)


